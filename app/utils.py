import pdfplumber
from langchain_openai import ChatOpenAI
from pdf2image import convert_from_path
import streamlit as st
import tempfile
import io
import base64
import logging
import pandas as pd
from docx import Document as DocxDocument
import magic
from PIL import Image
import os

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


def get_file_type(file_path):
    """ファイルの種類を判定する"""
    mime = magic.Magic(mime=True)
    return mime.from_file(file_path)

def extract_text_from_excel(file_path):
    """Excelファイルからテキストを抽出"""
    try:
        df = pd.read_excel(file_path)
        return df.to_string()
    except Exception as e:
        logger.error(f"Excelファイルの処理中にエラーが発生: {e}")
        return ""

def extract_text_from_word(file_path):
    """Wordファイルからテキストを抽出（テーブル、ヘッダー、フッター、画像内のテキストを含む）"""
    try:
        doc = DocxDocument(file_path)
        text_parts = []
        # 段落からテキストを抽出
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        # テーブルからテキストを抽出
        for table in doc.tables:
            for row in table.rows:
                row_texts = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_texts.append(cell.text.strip())
                if row_texts:
                    text_parts.append(" | ".join(row_texts))
        # ヘッダーからテキストを抽出
        for section in doc.sections:
            if section.header:
                for paragraph in section.header.paragraphs:
                    if paragraph.text.strip():
                        text_parts.append(f"[Header] {paragraph.text}")
        # フッターからテキストを抽出
        for section in doc.sections:
            if section.footer:
                for paragraph in section.footer.paragraphs:
                    if paragraph.text.strip():
                        text_parts.append(f"[Footer] {paragraph.text}")
        # 画像内のテキストを抽出（画像が存在する場合）
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                try:
                    image_data = rel.target_part.blob
                    with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp:
                        tmp.write(image_data)
                        tmp_path = tmp.name
                    image = Image.open(tmp_path)
                    image_text = convert_image_to_text(image)
                    if image_text.strip():
                        text_parts.append(f"[Image] {image_text}")
                    os.unlink(tmp_path)
                except Exception as e:
                    logger.error(f"画像内のテキスト抽出中にエラーが発生: {e}")
        # テキストを結合して返す
        extracted_text = "\n".join(text_parts)
        if not extracted_text.strip():
            logger.warning("Wordファイルからテキストを抽出できませんでした")
        return extracted_text

    except Exception as e:
        logger.error(f"Wordファイルの処理中にエラーが発生: {e}")
        return ""

def extract_text_from_image(file_path):
    """画像ファイルからテキストを抽出"""
    try:
        image = Image.open(file_path)
        # RGBA画像をRGBに変換
        if image.mode == 'RGBA':
            # 白背景を作成
            background = Image.new('RGB', image.size, (255, 255, 255))
            # RGBA画像をRGBに変換して白背景に合成
            background.paste(image, mask=image.split()[3])
            image = background
        return convert_image_to_text(image)
    except Exception as e:
        logger.error(f"画像ファイルの処理中にエラーが発生: {e}")
        return ""

def extract_text_from_pdf(pdf_path):
    with pdfplumber.open(pdf_path) as pdf:
        return "\n".join([page.extract_text() for page in pdf.pages])

def convert_pdf_to_images(pdf_path):
    # 一時ディレクトリを使ってPDF→画像変換
    with tempfile.TemporaryDirectory() as path:
        images = convert_from_path(pdf_path, output_folder=path, fmt='jpeg', dpi=200)
        return images

def convert_image_to_text(image):
    # Convert PIL Image to base64
    buf = io.BytesIO()
    image.save(buf, format="JPEG")
    encoded_image = base64.b64encode(buf.getvalue()).decode("utf-8")
    logger.info(f"✅ 画像をbase64に変換完了 : 画像は{encoded_image[:100]}")
    # llmを使って画像からテキストを抽出
    llm = ChatOpenAI(model_name="gpt-4o-mini", temperature=0)
    message = {
        "role": "user",
        "content": [
            {
                "type": "text",
                "text": st.session_state.custom_prompt
            },
            {
                "type": "image_url",
                "image_url": {
                    "url": f"data:image/jpeg;base64,{encoded_image}"
                }
            }
        ]
    }
    logger.info(f"✅ 画像からテキストを抽出中...")
    response = llm.invoke([message])
    logger.info(f"✅ 画像からテキストを抽出完了 : 回答 => {response.content[:100]}")
    return response.content

def process_file(file_path):
    """ファイルの種類に応じて適切な処理を行う"""
    file_type = get_file_type(file_path)
    if file_type == 'application/pdf':
        text = extract_text_from_pdf(file_path)
        if not text.strip():
            images = convert_pdf_to_images(file_path)
            text = "\n".join([convert_image_to_text(image) for image in images])
    elif file_type in ['application/vnd.openxmlformats-officedocument.spreadsheetml.sheet', 'application/vnd.ms-excel']:
        text = extract_text_from_excel(file_path)
    elif file_type in ['application/vnd.openxmlformats-officedocument.wordprocessingml.document', 'application/msword']:
        text = extract_text_from_word(file_path)
    elif file_type.startswith('image/'):
        text = extract_text_from_image(file_path)
    else:
        raise ValueError(f"未対応のファイル形式です: {file_type}")
    return text
